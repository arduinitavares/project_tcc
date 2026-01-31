import docx
import sys
from pathlib import Path

def convert_to_md(docx_file: Path) -> Path:
    if not docx_file.exists():
        print(f"Error: File {docx_file} not found.")
        sys.exit(1)

    print(f"Reading {docx_file}...")
    try:
        doc = docx.Document(docx_file)
    except Exception as e:
        print(f"Failed to open document: {e}")
        sys.exit(1)

    md_lines = []
    
    # Extract paragraphs
    for para in doc.paragraphs:
        # Simple style mapping could be added here
        text = para.text.strip()
        if not text:
            continue
            
        style_name = para.style.name.lower()
        if 'heading 1' in style_name:
            md_lines.append(f"# {text}")
        elif 'heading 2' in style_name:
            md_lines.append(f"## {text}")
        elif 'heading 3' in style_name:
            md_lines.append(f"### {text}")
        elif 'list bullet' in style_name:
            md_lines.append(f"- {text}")
        elif 'list number' in style_name:
            md_lines.append(f"1. {text}") # Simple numbering
        else:
            md_lines.append(text)
            
    # Extract tables (basic)
    if doc.tables:
        md_lines.append("\n\n--- TABLES ---\n")
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                md_lines.append(" | ".join(row_text))
            md_lines.append("") # Empty line between tables

    output_file = docx_file.with_suffix('.md')
    
    print(f"Writing to {output_file}...")
    with open(output_file, 'w', encoding='utf-8') as f:
        f.write('\n\n'.join(md_lines))
        
    print("Done.")
    return output_file

if __name__ == "__main__":
    docx_path = Path("Proposta_TCC_Alexandre_Tavares.docx").resolve()
    convert_to_md(docx_path)
